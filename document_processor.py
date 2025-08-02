#!/usr/bin/env python3
"""
Advanced Legal Document Processing System
Handles document upload, parsing, and legal domain classification
"""

import os
import sys
import json
import uuid
import hashlib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import logging

# Document processing libraries
try:
    import fitz  # PyMuPDF
    import pytesseract
    from PIL import Image
    import docx
    import pandas as pd
    from textstat import flesch_reading_ease, flesch_kincaid_grade
except ImportError as e:
    print(f"⚠️ Missing dependencies: {e}")
    print("📦 Install with: pip install PyMuPDF pytesseract pillow python-docx pandas textstat")

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LegalDocumentProcessor:
    """Advanced legal document processor with AI-powered classification"""
    
    def __init__(self, upload_dir: str = "uploads", processed_dir: str = "processed"):
        self.upload_dir = Path(upload_dir)
        self.processed_dir = Path(processed_dir)
        self.upload_dir.mkdir(exist_ok=True)
        self.processed_dir.mkdir(exist_ok=True)
        
        # Legal document types and their keywords
        self.legal_document_types = {
            "rental_agreement": {
                "keywords": ["rent", "lease", "tenant", "landlord", "property", "monthly", "deposit", "premises"],
                "patterns": ["lease agreement", "rental contract", "tenancy agreement"],
                "priority_sections": ["rent amount", "lease term", "security deposit", "termination"]
            },
            "employment_contract": {
                "keywords": ["employment", "employee", "employer", "salary", "wages", "benefits", "termination"],
                "patterns": ["employment agreement", "job contract", "work agreement"],
                "priority_sections": ["compensation", "job duties", "termination clause", "benefits"]
            },
            "legal_notice": {
                "keywords": ["notice", "demand", "cease", "desist", "violation", "breach", "legal action"],
                "patterns": ["legal notice", "demand letter", "cease and desist"],
                "priority_sections": ["violation details", "demanded action", "deadline", "consequences"]
            },
            "contract": {
                "keywords": ["contract", "agreement", "party", "obligations", "terms", "conditions"],
                "patterns": ["service agreement", "purchase agreement", "contract"],
                "priority_sections": ["parties", "obligations", "payment terms", "termination"]
            },
            "court_document": {
                "keywords": ["court", "plaintiff", "defendant", "case", "docket", "hearing", "judgment"],
                "patterns": ["court order", "summons", "complaint", "motion"],
                "priority_sections": ["case number", "parties", "hearing date", "orders"]
            },
            "will_testament": {
                "keywords": ["will", "testament", "executor", "beneficiary", "estate", "inherit"],
                "patterns": ["last will", "testament", "estate planning"],
                "priority_sections": ["executor", "beneficiaries", "assets", "instructions"]
            }
        }
        
        # Supported file types
        self.supported_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.rtf': 'application/rtf',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg'
        }

    def validate_file(self, file_path: Path) -> Dict[str, Any]:
        """Validate uploaded file"""
        try:
            if not file_path.exists():
                return {"valid": False, "error": "File does not exist"}
            
            # Check file size (max 50MB)
            file_size = file_path.stat().st_size
            if file_size > 50 * 1024 * 1024:
                return {"valid": False, "error": "File too large (max 50MB)"}
            
            # Check file type
            file_ext = file_path.suffix.lower()
            if file_ext not in self.supported_types:
                return {"valid": False, "error": f"Unsupported file type: {file_ext}"}
            
            # Generate file hash for deduplication
            file_hash = self._calculate_file_hash(file_path)
            
            return {
                "valid": True,
                "file_size": file_size,
                "file_type": file_ext,
                "mime_type": self.supported_types[file_ext],
                "file_hash": file_hash
            }
            
        except Exception as e:
            logger.error(f"File validation error: {e}")
            return {"valid": False, "error": str(e)}

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def extract_text_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text_content = []
            metadata = {}
            
            # Extract metadata
            metadata = {
                "page_count": doc.page_count,
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", "")
            }
            
            # Extract text from each page
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                
                # Extract images if text is minimal (OCR candidate)
                if len(text.strip()) < 50:
                    image_list = page.get_images()
                    if image_list:
                        # Perform OCR on images
                        ocr_text = self._perform_ocr_on_page(page)
                        text += f"\n[OCR Content]: {ocr_text}"
                
                text_content.append({
                    "page": page_num + 1,
                    "text": text,
                    "char_count": len(text)
                })
            
            doc.close()
            
            # Combine all text
            full_text = "\n".join([page["text"] for page in text_content])
            
            return {
                "success": True,
                "text": full_text,
                "pages": text_content,
                "metadata": metadata,
                "extraction_method": "PyMuPDF"
            }
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return {"success": False, "error": str(e)}

    def _perform_ocr_on_page(self, page) -> str:
        """Perform OCR on PDF page images"""
        try:
            # Get page as image
            mat = fitz.Matrix(2, 2)  # Increase resolution
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Convert to PIL Image
            from io import BytesIO
            img = Image.open(BytesIO(img_data))
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(img)
            return ocr_text
            
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return ""

    def extract_text_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables_text.append("\n".join(table_data))
            
            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n[TABLES]\n" + "\n\n".join(tables_text)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
            
            return {
                "success": True,
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables_text,
                "metadata": metadata,
                "extraction_method": "python-docx"
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return {"success": False, "error": str(e)}

    def extract_text_from_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from images using OCR"""
        try:
            img = Image.open(file_path)
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(img)
            
            # Get image metadata
            metadata = {
                "format": img.format,
                "mode": img.mode,
                "size": img.size,
                "has_transparency": img.mode in ("RGBA", "LA") or "transparency" in img.info
            }
            
            return {
                "success": True,
                "text": ocr_text,
                "metadata": metadata,
                "extraction_method": "Tesseract OCR"
            }
            
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            return {"success": False, "error": str(e)}

    def extract_text(self, file_path: Path) -> Dict[str, Any]:
        """Main text extraction method"""
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx']:
            return self.extract_text_from_docx(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return self.extract_text_from_image(file_path)
        elif file_ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                return {
                    "success": True,
                    "text": text,
                    "metadata": {"char_count": len(text)},
                    "extraction_method": "Direct text read"
                }
            except Exception as e:
                return {"success": False, "error": str(e)}
        else:
            return {"success": False, "error": f"Unsupported file type: {file_ext}"}

    def classify_document(self, text: str) -> Dict[str, Any]:
        """Classify document type based on content analysis"""
        text_lower = text.lower()
        classification_scores = {}

        # Score each document type
        for doc_type, config in self.legal_document_types.items():
            score = 0
            matched_keywords = []
            matched_patterns = []

            # Check keywords
            for keyword in config["keywords"]:
                count = text_lower.count(keyword.lower())
                if count > 0:
                    score += count * 2
                    matched_keywords.append(keyword)

            # Check patterns (higher weight)
            for pattern in config["patterns"]:
                if pattern.lower() in text_lower:
                    score += 10
                    matched_patterns.append(pattern)

            classification_scores[doc_type] = {
                "score": score,
                "matched_keywords": matched_keywords,
                "matched_patterns": matched_patterns
            }

        # Find best match
        best_match = max(classification_scores.items(), key=lambda x: x[1]["score"])

        # Calculate confidence
        total_score = sum(item["score"] for item in classification_scores.values())
        confidence = (best_match[1]["score"] / total_score * 100) if total_score > 0 else 0

        return {
            "document_type": best_match[0],
            "confidence": round(confidence, 2),
            "all_scores": classification_scores,
            "is_legal_document": confidence > 20
        }

    def extract_key_information(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract key information based on document type"""
        key_info = {
            "entities": [],
            "dates": [],
            "amounts": [],
            "parties": [],
            "key_sections": {}
        }

        # Extract dates
        import re
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{2,4}\b',
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4}\b'
        ]

        for pattern in date_patterns:
            dates = re.findall(pattern, text, re.IGNORECASE)
            key_info["dates"].extend(dates)

        # Extract monetary amounts
        money_patterns = [
            r'\$[\d,]+\.?\d*',
            r'\b\d+\s*dollars?\b',
            r'\b\d+\s*USD\b'
        ]

        for pattern in money_patterns:
            amounts = re.findall(pattern, text, re.IGNORECASE)
            key_info["amounts"].extend(amounts)

        # Extract parties (names in quotes or after specific keywords)
        party_patterns = [
            r'"([^"]+)"',
            r'between\s+([^,\n]+)\s+and\s+([^,\n]+)',
            r'tenant[:\s]+([^,\n]+)',
            r'landlord[:\s]+([^,\n]+)',
            r'employee[:\s]+([^,\n]+)',
            r'employer[:\s]+([^,\n]+)'
        ]

        for pattern in party_patterns:
            parties = re.findall(pattern, text, re.IGNORECASE)
            if isinstance(parties[0], tuple) if parties else False:
                key_info["parties"].extend([p for party_tuple in parties for p in party_tuple])
            else:
                key_info["parties"].extend(parties)

        # Extract key sections based on document type
        if document_type in self.legal_document_types:
            priority_sections = self.legal_document_types[document_type]["priority_sections"]
            for section in priority_sections:
                section_content = self._extract_section_content(text, section)
                if section_content:
                    key_info["key_sections"][section] = section_content

        return key_info

    def _extract_section_content(self, text: str, section_name: str) -> str:
        """Extract content from a specific section"""
        import re

        # Create pattern to find section
        pattern = rf'{re.escape(section_name)}[:\s]*([^\n]*(?:\n(?!\s*[A-Z][^:]*:)[^\n]*)*)'
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)

        if match:
            return match.group(1).strip()
        return ""

    def analyze_document_complexity(self, text: str) -> Dict[str, Any]:
        """Analyze document complexity and readability"""
        try:
            # Basic statistics
            word_count = len(text.split())
            sentence_count = text.count('.') + text.count('!') + text.count('?')
            paragraph_count = text.count('\n\n') + 1

            # Readability scores
            flesch_score = flesch_reading_ease(text)
            fk_grade = flesch_kincaid_grade(text)

            # Legal complexity indicators
            legal_terms = [
                'whereas', 'heretofore', 'hereinafter', 'notwithstanding', 'pursuant',
                'covenant', 'indemnify', 'liability', 'jurisdiction', 'arbitration'
            ]

            legal_term_count = sum(text.lower().count(term) for term in legal_terms)

            return {
                "word_count": word_count,
                "sentence_count": sentence_count,
                "paragraph_count": paragraph_count,
                "avg_words_per_sentence": round(word_count / sentence_count, 2) if sentence_count > 0 else 0,
                "flesch_reading_ease": round(flesch_score, 2),
                "flesch_kincaid_grade": round(fk_grade, 2),
                "legal_term_density": round(legal_term_count / word_count * 100, 2) if word_count > 0 else 0,
                "complexity_level": self._determine_complexity_level(flesch_score, legal_term_count, word_count)
            }

        except Exception as e:
            logger.error(f"Complexity analysis error: {e}")
            return {"error": str(e)}

    def _determine_complexity_level(self, flesch_score: float, legal_terms: int, word_count: int) -> str:
        """Determine document complexity level"""
        legal_density = (legal_terms / word_count * 100) if word_count > 0 else 0

        if flesch_score >= 60 and legal_density < 2:
            return "Simple"
        elif flesch_score >= 30 and legal_density < 5:
            return "Moderate"
        elif flesch_score >= 0 and legal_density < 10:
            return "Complex"
        else:
            return "Very Complex"

    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """Main document processing pipeline"""
        try:
            # Generate unique processing ID
            process_id = str(uuid.uuid4())

            # Validate file
            validation = self.validate_file(file_path)
            if not validation["valid"]:
                return {"success": False, "error": validation["error"]}

            # Extract text
            extraction_result = self.extract_text(file_path)
            if not extraction_result["success"]:
                return {"success": False, "error": extraction_result["error"]}

            text = extraction_result["text"]

            # Classify document
            classification = self.classify_document(text)

            # Extract key information
            key_info = self.extract_key_information(text, classification["document_type"])

            # Analyze complexity
            complexity = self.analyze_document_complexity(text)

            # Create processing result
            result = {
                "success": True,
                "process_id": process_id,
                "timestamp": datetime.now().isoformat(),
                "file_info": {
                    "name": file_path.name,
                    "size": validation["file_size"],
                    "type": validation["file_type"],
                    "hash": validation["file_hash"]
                },
                "extraction": extraction_result,
                "classification": classification,
                "key_information": key_info,
                "complexity_analysis": complexity,
                "text_preview": text[:500] + "..." if len(text) > 500 else text
            }

            # Save processing result
            result_file = self.processed_dir / f"{process_id}.json"
            with open(result_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)

            logger.info(f"Document processed successfully: {process_id}")
            return result

        except Exception as e:
            logger.error(f"Document processing error: {e}")
            return {"success": False, "error": str(e)}

    def get_processing_summary(self, process_id: str) -> Dict[str, Any]:
        """Get summary of processed document"""
        try:
            result_file = self.processed_dir / f"{process_id}.json"
            if not result_file.exists():
                return {"success": False, "error": "Processing result not found"}

            with open(result_file, 'r', encoding='utf-8') as f:
                result = json.load(f)

            # Create summary
            summary = {
                "process_id": process_id,
                "document_type": result["classification"]["document_type"],
                "confidence": result["classification"]["confidence"],
                "complexity": result["complexity_analysis"]["complexity_level"],
                "key_parties": result["key_information"]["parties"][:3],  # Top 3
                "important_dates": result["key_information"]["dates"][:3],  # Top 3
                "monetary_amounts": result["key_information"]["amounts"][:3],  # Top 3
                "word_count": result["complexity_analysis"]["word_count"],
                "processing_time": result["timestamp"]
            }

            return {"success": True, "summary": summary}

        except Exception as e:
            logger.error(f"Summary generation error: {e}")
            return {"success": False, "error": str(e)}
